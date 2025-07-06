import os
import random
import re
from docx import Document
from docx.shared import Inches, Pt  # Pt pentru dimensiunea fontului
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Pentru alinierea textului
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.platypus import Image as ReportLabImage


QUIZ_BASE_DIR = "final"  # Directorul rădăcină care conține "bio" și "chimie"
# Directorul unde vor fi salvate simulările
OUTPUT_SIMULATION_DIR = "simulari_generate"

# Criterii de selecție
NUM_BIO_QUIZZES = 35
NUM_CHIMIE_TEORIE_QUIZZES = 10
NUM_CHIMIE_PROBLEME_QUIZZES = 5

# bio
# cap1_corpul_uman_celula
# cap2_oasele_articulatiile
# cap3_tesuturi_excitabile
# cap4_sistemul_nervos
# cap5_organe_de_simt
# cap6_sistemul_endocrin_metabolism
# cap7_sangele
# cap8_sistemul_circulator
# cap9_sistemul_respirator
# cap10_sistemul_digestiv
# cap11_sistemul_urinar
# cap12_sistemul_reproducator
# cap13_intrebari_asociative_recap
BIOLOGY_CHAPTER_WEIGHTS = {
    1: 0,
    2: 0,
    3: 0,
    4: 0,
    5: 0,
    6: 0,
    7: 0,
    8: 2,
    9: 0,
    10: 0,
    11: 1,
    12: 5,
    13: 0.5,

    # Dacă un capitol nu este listat, va avea ponderea implicită 1.0.
    # 1 - basic, 0.5 - sanse la jumatate, 2 - sanse duble
}

# ----------------------------------------

# chimie
# cap1_solutii_acizi_baze
# cap2_compozitia_structura_compusilor_organici
# cap3_compusi_hidroxilici
# cap4_amine
# cap5_aldehide_cetone
# cap6_acizi_carboxilici
# cap7_proteine
# cap8_glucide
# cap9_medicamente_droguri
# cap10_izomerie
# cap11_grile_asociative_recap

CHEMISTRY_CHAPTER_WEIGHTS_TEORIE = {
    1: 1,
    2: 1,
    3: 1,
    4: 1,
    5: 1,
    6: 1,
    7: 1,
    8: 1,
    9: 1,
    10: 1,
    11: 1,
}

CHEMISTRY_CHAPTER_WEIGHTS_PROBLEME = {
    1: 1,
    2: 1,
    3: 1,
    4: 1,
    5: 1,
    6: 1,
    7: 1,
    8: 1,
    9: 1,
    10: 1,
    11: 1,
}


def get_quizzes_by_subject_and_type(base_dir):
    """
    Parcurge directoarele și colectează căile către grile,
    clasificându-le după subiect și tip (teorie/probleme), conform noii structuri.
    """
    quizzes = {
        # Vom pune toate grilele bio în "teorie"
        "biologie": {"teorie": [], "probleme": []},
        "chimie": {"teorie": [], "probleme": []},
    }

    print(f"Colectez grile din: {base_dir}")

    for subject_folder_name in os.listdir(base_dir):
        subject_path = os.path.join(base_dir, subject_folder_name)

        if not os.path.isdir(subject_path):
            continue

        current_subject = None
        if "bio" in subject_folder_name.lower():
            current_subject = "biologie"
        elif "chimie" in subject_folder_name.lower():
            current_subject = "chimie"
        else:
            print(
                f"Avertisment: Director subiect necunoscut '{subject_folder_name}'. Ignor.")
            continue

        print(f"  Procesez subiectul: {current_subject}")

        for chapter_folder_name in os.listdir(subject_path):
            chapter_folder_path = os.path.join(
                subject_path, chapter_folder_name)

            if not os.path.isdir(chapter_folder_path):
                continue

            match_chapter = re.match(r"cap(\d+)_.*", chapter_folder_name)
            if not match_chapter:
                continue

            chapter_number = match_chapter.group(1)

            done_folder_path = os.path.join(
                chapter_folder_path, "grile", "done")

            if not os.path.exists(done_folder_path):
                continue

            # Colectează grilele de TEORIE
            for quiz_filename in os.listdir(done_folder_path):
                if quiz_filename.lower().endswith(".png"):
                    quiz_path = os.path.join(done_folder_path, quiz_filename)

                    quiz_number_match = re.match(
                        r"(\d+)\.png", quiz_filename, re.IGNORECASE)
                    if quiz_number_match:
                        quiz_number = int(quiz_number_match.group(1))
                    else:
                        print(
                            f"    Avertisment: Nume fișier grilă neașteptat '{quiz_filename}'. Sărit.")
                        continue

                    quiz_info = {
                        "path": quiz_path,
                        "quiz_number": quiz_number,
                        "chapter_number": int(chapter_number)
                    }

                    quizzes[current_subject]["teorie"].append(quiz_info)

            # Colectează grilele de PROBLEME (doar pentru chimie)
            if current_subject == "chimie":
                pb_folder_path = os.path.join(done_folder_path, "pb")
                if os.path.exists(pb_folder_path) and os.path.isdir(pb_folder_path):
                    for quiz_filename in os.listdir(pb_folder_path):
                        if quiz_filename.lower().endswith(".png"):
                            quiz_path = os.path.join(
                                pb_folder_path, quiz_filename)

                            quiz_number_match = re.match(
                                r"(\d+)\.png", quiz_filename, re.IGNORECASE)
                            if quiz_number_match:
                                quiz_number = int(quiz_number_match.group(1))
                            else:
                                print(
                                    f"    Avertisment: Nume fișier grilă problemă neașteptat '{quiz_filename}'. Sărit.")
                                continue

                            quiz_info = {
                                "path": quiz_path,
                                "quiz_number": quiz_number,
                                "chapter_number": int(chapter_number)
                            }
                            quizzes[current_subject]["probleme"].append(
                                quiz_info)

    print("Colectare grile finalizată.")

    print(f"Sumar grile colectate:")
    print(f"  Biologie (total): {len(quizzes['biologie']['teorie'])}")
    print(f"  Chimie (teorie): {len(quizzes['chimie']['teorie'])}")
    print(f"  Chimie (probleme): {len(quizzes['chimie']['probleme'])}")

    return quizzes


def create_word_document(selected_bio_quizzes, selected_chimie_quizzes, output_path, chapter_quiz_map_lines):
    """Generează un document Word cu imaginile selectate, ordonate după subiect."""
    document = Document()

    for i, quiz_info in enumerate(selected_bio_quizzes):
        try:
            document.add_picture(quiz_info["path"], width=Inches(6))
        except Exception as e:
            document.add_paragraph(
                f"Eroare la adăugarea imaginii {quiz_info['path']}: {e}")

    # Adaugă delimitatorul '-'
    delimiter_paragraph = document.add_paragraph('-')
    delimiter_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrează delimitatorul
    delimiter_paragraph.runs[0].font.size = Pt(
        24)  # Mărește fontul pentru vizibilitate
    document.add_page_break()  # Asigură-te că secțiunea de chimie începe pe o pagină nouă

    start_num_chimie = len(selected_bio_quizzes)
    for i, quiz_info in enumerate(selected_chimie_quizzes):
        try:
            document.add_picture(quiz_info["path"], width=Inches(6))
        except Exception as e:
            document.add_paragraph(
                f"Eroare la adăugarea imaginii {quiz_info['path']}: {e}")

    document.save(output_path)
    print(f"Document Word generat: {output_path}")


def create_pdf_document(selected_bio_quizzes, selected_chimie_quizzes, output_path, chapter_quiz_map_lines):
    """
    Generează un document PDF cu imaginile selectate, respectând ordinea și delimitarea
    din documentul Word.
    """
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4

    # Adaugă informațiile din fișierul TXT la începutul documentului PDF
    textobject = c.beginText()
    textobject.setTextOrigin(inch, height - inch)
    textobject.setFont("Helvetica", 12)
    # c.drawText(textobject)

    # Poziție de start pentru imagini, sub textul inițial
    y_position = height - (2 * inch)
    margin = 0.5 * inch

    # Funcție ajutătoare pentru adăugarea unei grile
    def add_quiz_to_pdf(quiz_info, current_y, quiz_idx):
        nonlocal c, y_position  # 'nonlocal' pentru a modifica variabilele din scope-ul exterior
        try:
            img_path = quiz_info["path"]
            temp_img = ReportLabImage(img_path)
            img_width, img_height = temp_img.drawWidth, temp_img.drawHeight

            available_width = width - 2 * margin
            if img_width > available_width:
                scale_factor = available_width / img_width
                img_width *= scale_factor
                img_height *= scale_factor

            # Adaugă numărul grilei deasupra imaginii
            c.setFont("Helvetica-Bold", 10)
            current_y -= 0.3 * inch  # Spațiu pentru numărul grilei

            if current_y - img_height < margin:  # Verifică dacă imaginea încape pe pagina curentă
                c.showPage()
                y_position = height - inch  # Resetează poziția Y pe pagina nouă
                current_y = y_position
                # Resetează fontul după schimbarea paginii
                c.setFont("Helvetica-Bold", 10)
                current_y -= 0.3 * inch

            c.drawImage(img_path, margin, current_y - img_height,
                        width=img_width, height=img_height)
            current_y -= (img_height + 0.2 * inch)  # Spațiu între imagini

        except Exception as e:
            print(
                f"Eroare la adăugarea imaginii {quiz_info['path']} în PDF: {e}")
            c.drawString(margin, current_y,
                         f"Eroare la încărcarea imaginii: {quiz_info['path']}")
            current_y -= (0.5 * inch)
        return current_y

    # Adaugă grilele de Biologie
    c.setFont("Helvetica-Bold", 14)
    y_position -= 0.5 * inch  # Spațiu după titlu
    for i, quiz_info in enumerate(selected_bio_quizzes):
        y_position = add_quiz_to_pdf(quiz_info, y_position, i + 1)

    # Adaugă delimitatorul '-'
    if y_position - 0.5 * inch < margin:  # Verifică dacă delimitatorul încape
        c.showPage()
        y_position = height - inch
    c.setFont("Helvetica-Bold", 24)
    c.drawCentredString(width / 2, y_position - 0.3 * inch, "-")
    y_position -= 0.5 * inch  # Spațiu după delimitator
    c.showPage()  # Trece la o pagină nouă pentru chimie

    # Adaugă grilele de Chimie
    c.setFont("Helvetica-Bold", 14)
    y_position -= 0.5 * inch  # Spațiu după titlu
    start_num_chimie = len(selected_bio_quizzes)
    for i, quiz_info in enumerate(selected_chimie_quizzes):
        y_position = add_quiz_to_pdf(
            quiz_info, y_position, start_num_chimie + i + 1)

    c.save()
    print(f"Document PDF generat: {output_path}")


def select_quizzes_with_weights(all_quizzes_list, num_needed, weights):
    """
    Selectează un număr specific de grile dintr-o listă generală,
    ținând cont de ponderile specificate pentru fiecare capitol.

    Args:
        all_quizzes_list (list): O listă de dicționare quiz_info din care să selectezi.
        num_needed (int): Numărul total de grile de selectat.
        weights (dict): Dicționar de ponderi {chapter_number: weight}.
                        Capitolele fără pondere explicită vor avea ponderea 1.0.

    Returns:
        list: O listă de quiz_info selectate, unice și ponderate.
    """
    # Grupează grilele pe capitole
    quizzes_by_chapter = {}
    for quiz in all_quizzes_list:
        chapter_num = quiz["chapter_number"]
        if chapter_num not in quizzes_by_chapter:
            quizzes_by_chapter[chapter_num] = []
        quizzes_by_chapter[chapter_num].append(quiz)

    selection_pool_chapters = []  # Lista de capitole din care vom extrage aleatoriu

    # Construiește "pool-ul" de selecție a capitolelor, ponderat
    for chapter_num, quizzes_in_chapter in quizzes_by_chapter.items():
        weight = weights.get(chapter_num, 1.0)
        # Adaugă capitolul de `weight * 10` ori (sau un factor similar) în pool.
        # Factorul 10 este pentru a da o granulație suficientă la ponderi precum 0.5.
        if quizzes_in_chapter:  # Asigură-te că sunt grile în capitol
            for _ in range(int(weight * 10)):
                selection_pool_chapters.append(chapter_num)

    if not selection_pool_chapters:
        print(
            "Avertisment: Pool-ul de selecție capitole este gol. Nu se pot extrage grile.")
        return []

    selected_quizzes_unique = []
    selected_quiz_paths = set()  # Pentru a asigura unicitatea grilelor selectate

    attempts = 0
    max_attempts = num_needed * 10 + 200  # Creștem numărul maxim de încercări

    while len(selected_quizzes_unique) < num_needed and attempts < max_attempts:
        attempts += 1
        # Alege un capitol din pool-ul ponderat
        chosen_chapter_num = random.choice(selection_pool_chapters)

        quizzes_from_chosen_chapter = quizzes_by_chapter.get(
            chosen_chapter_num, [])

        # Filtrăm grilele deja selectate din capitolul curent pentru a nu le alege din nou
        available_in_chapter = [
            q for q in quizzes_from_chosen_chapter if q["path"] not in selected_quiz_paths
        ]

        if not available_in_chapter:
            # Nu mai sunt grile unice disponibile în acest capitol.
            # Încercăm să eliminăm acest capitol din pool pentru a nu-l mai alege.
            # Aceasta este o operație costisitoare, dar necesară dacă pool-ul e mare.
            selection_pool_chapters = [
                c for c in selection_pool_chapters if c != chosen_chapter_num
            ]
            if not selection_pool_chapters:
                # Nu mai sunt capitole din care să extragem grile unice
                print(
                    f"Avertisment: Nu mai sunt grile unice disponibile în niciun capitol pentru a atinge numărul necesar ({num_needed}). S-au găsit doar {len(selected_quizzes_unique)}.")
                break  # Ieșim din bucla while
            continue

        # Alege o grilă aleatorie din cele rămase unice din capitolul ales
        quiz_candidate = random.choice(available_in_chapter)

        selected_quizzes_unique.append(quiz_candidate)
        selected_quiz_paths.add(quiz_candidate["path"])

    if len(selected_quizzes_unique) < num_needed:
        print(
            f"Avertisment: Nu s-au putut selecta {num_needed} grile unice cu ponderile și grilele disponibile. S-au găsit doar {len(selected_quizzes_unique)}.")

    return selected_quizzes_unique


def generate_simulation():
    """Generează o simulare conform specificațiilor."""
    all_quizzes = get_quizzes_by_subject_and_type(QUIZ_BASE_DIR)

    bio_quizzes = all_quizzes["biologie"]["teorie"]
    chimie_teorie_quizzes = all_quizzes["chimie"]["teorie"]
    chimie_probleme_quizzes = all_quizzes["chimie"]["probleme"]

    # --- Validări inițiale (rămân la fel) ---
    if len(bio_quizzes) < NUM_BIO_QUIZZES:
        print(
            f"Eroare: Nu sunt suficiente grile de biologie ({len(bio_quizzes)} disponibile, necesare {NUM_BIO_QUIZZES}).")
        return
    if len(chimie_teorie_quizzes) < NUM_CHIMIE_TEORIE_QUIZZES:
        print(
            f"Eroare: Nu sunt suficiente grile de chimie (teorie) ({len(chimie_teorie_quizzes)} disponibile, necesare {NUM_CHIMIE_TEORIE_QUIZZES}).")
        return
    if len(chimie_probleme_quizzes) < NUM_CHIMIE_PROBLEME_QUIZZES:
        print(
            f"Eroare: Nu sunt suficiente grile de chimie (probleme) ({len(chimie_probleme_quizzes)} disponibile, necesare {NUM_CHIMIE_PROBLEME_QUIZZES}).")
        return

    # --- Selecția grilelor folosind ponderi ---
    print(f"Selecționez {NUM_BIO_QUIZZES} grile de biologie cu ponderi...")
    selected_bio = select_quizzes_with_weights(
        bio_quizzes, NUM_BIO_QUIZZES, BIOLOGY_CHAPTER_WEIGHTS
    )
    if len(selected_bio) < NUM_BIO_QUIZZES:
        print(
            f"ATENȚIE: Nu s-au putut selecta toate grilele de biologie dorite cu ponderi. S-au obținut doar {len(selected_bio)}.")

    print(
        f"Selecționez {NUM_CHIMIE_TEORIE_QUIZZES} grile de chimie (teorie) cu ponderi...")
    selected_chimie_teorie = select_quizzes_with_weights(
        chimie_teorie_quizzes, NUM_CHIMIE_TEORIE_QUIZZES, CHEMISTRY_CHAPTER_WEIGHTS_TEORIE
    )
    if len(selected_chimie_teorie) < NUM_CHIMIE_TEORIE_QUIZZES:
        print(
            f"ATENȚIE: Nu s-au putut selecta toate grilele de chimie (teorie) dorite cu ponderi. S-au obținut doar {len(selected_chimie_teorie)}.")

    print(
        f"Selecționez {NUM_CHIMIE_PROBLEME_QUIZZES} grile de chimie (probleme) cu ponderi...")
    selected_chimie_probleme = select_quizzes_with_weights(
        chimie_probleme_quizzes, NUM_CHIMIE_PROBLEME_QUIZZES, CHEMISTRY_CHAPTER_WEIGHTS_PROBLEME
    )
    if len(selected_chimie_probleme) < NUM_CHIMIE_PROBLEME_QUIZZES:
        print(
            f"ATENȚIE: Nu s-au putut selecta toate grilele de chimie (probleme) dorite cu ponderi. S-au obținut doar {len(selected_chimie_probleme)}.")

    selected_bio_quizzes = selected_bio
    selected_chimie_combined = selected_chimie_teorie + selected_chimie_probleme
    # Amestecă grilele de chimie (teorie și probleme)
    random.shuffle(selected_chimie_combined)

    # --- Generare fișier TXT cu detalii (rămâne la fel) ---
    bio_map = {}
    for quiz in selected_bio_quizzes:
        chapter = quiz["chapter_number"]
        quiz_num = quiz["quiz_number"]
        if chapter not in bio_map:
            bio_map[chapter] = []
        bio_map[chapter].append(quiz_num)

    chimie_map = {}
    for quiz in selected_chimie_combined:
        chapter = quiz["chapter_number"]
        quiz_num = quiz["quiz_number"]
        if chapter not in chimie_map:
            chimie_map[chapter] = []
        chimie_map[chapter].append(quiz_num)

    chapter_quiz_map_lines = []
    chapter_quiz_map_lines.append("BIO")
    for chapter in sorted(bio_map.keys()):
        quizzes_in_chapter = sorted(bio_map[chapter])
        chapter_quiz_map_lines.append(
            f"cap{chapter} -> grila {', '.join(map(str, quizzes_in_chapter))}")

    chapter_quiz_map_lines.append("")
    chapter_quiz_map_lines.append("CHIMIE")
    for chapter in sorted(chimie_map.keys()):
        quizzes_in_chapter = sorted(chimie_map[chapter])
        chapter_quiz_map_lines.append(
            f"cap{chapter} -> grila {', '.join(map(str, quizzes_in_chapter))}")

    os.makedirs(OUTPUT_SIMULATION_DIR, exist_ok=True)

    simulation_name = f"simulare_{random.randint(1000,9999)}"

    txt_output_path = os.path.join(
        OUTPUT_SIMULATION_DIR, f"{simulation_name}_detalii.txt")
    with open(txt_output_path, "w", encoding='utf-8') as f:
        f.write("Detalii Simulăre:\n")
        for line in chapter_quiz_map_lines:
            f.write(line + "\n")
    print(f"Fișier TXT generat: {txt_output_path}")

    # --- Generare documente Word și PDF (rămâne la fel) ---
    word_output_path = os.path.join(
        OUTPUT_SIMULATION_DIR, f"{simulation_name}.docx")
    create_word_document(selected_bio_quizzes, selected_chimie_combined,
                         word_output_path, chapter_quiz_map_lines)

    pdf_output_path = os.path.join(
        OUTPUT_SIMULATION_DIR, f"{simulation_name}.pdf")
    create_pdf_document(selected_bio_quizzes, selected_chimie_combined,
                        pdf_output_path, chapter_quiz_map_lines)

    print("\nGenerarea simulării finalizată!")


# --- Rularea Procesului Principal ---
if __name__ == "__main__":
    if not os.path.exists(QUIZ_BASE_DIR):
        print(
            f"Eroare: Directorul de bază pentru grile '{QUIZ_BASE_DIR}' nu există.")
        print("Asigură-te că structura ta este 'final/bio/...' și 'final/chimie/...'.")
    else:
        generate_simulation()
